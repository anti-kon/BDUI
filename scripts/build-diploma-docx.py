from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "vkr-part-2.md"
DEFAULT_OUT = ROOT / "reports" / "vkr-part-2-bdui-konopelkin-2026.docx"
GRAPH_OUT = ROOT / "docs" / "assets" / "bdui-package-graph.png"


def set_run_font(run, name: str, size_pt: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, *, before=0, after=0, line=1.5, first_indent=True):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    fmt.first_line_indent = Cm(1.25) if first_indent else None


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)

    for style_name, size, before, after in [
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 14, 10, 4),
        ("Heading 3", 14, 8, 3),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = None

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(0)

    code_style = doc.styles.add_style("Diploma Code", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Cm(0.5)
    code_style.paragraph_format.right_indent = Cm(0.2)
    code_style.paragraph_format.first_line_indent = None
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0

    caption_style = doc.styles.add_style("Diploma Caption", 1)
    caption_style.font.name = "Times New Roman"
    caption_style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption_style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    caption_style.font.size = Pt(12)
    caption_style.font.italic = True
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.first_line_indent = None
    caption_style.paragraph_format.space_before = Pt(3)
    caption_style.paragraph_format.space_after = Pt(6)
    caption_style.paragraph_format.line_spacing = 1.15


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_field(paragraph, instruction: str, placeholder: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)

    if placeholder:
        paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_page_number_footer(doc: Document):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE", "1")
    for run in paragraph.runs:
        set_run_font(run, "Times New Roman", 12)


def add_rich_text(paragraph, text: str, *, size=14, bold=False):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", size - 1, bold)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Times New Roman", size, bold)


def add_body_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph)
    add_rich_text(paragraph, text)
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Diploma Caption")
    add_rich_text(paragraph, text, size=12)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    add_rich_text(paragraph, text, bold=True)
    return paragraph


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        paragraph = doc.add_paragraph(style="Diploma Code")
        set_paragraph_shading(paragraph, "F4F6F8")
        run = paragraph.add_run(line if line else " ")
        set_run_font(run, "Consolas", 8.5)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str):
    draw.line([start, end], fill=fill, width=2)
    x1, y1 = start
    x2, y2 = end
    dx = x2 - x1
    dy = y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    left = (-uy, ux)
    size = 9
    point = (x2, y2)
    p1 = (x2 - ux * size + left[0] * size * 0.55, y2 - uy * size + left[1] * size * 0.55)
    p2 = (x2 - ux * size - left[0] * size * 0.55, y2 - uy * size - left[1] * size * 0.55)
    draw.polygon([point, p1, p2], fill=fill)


def create_dependency_graph(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 26)
        small = ImageFont.truetype("arial.ttf", 22)
    except OSError:
        font = ImageFont.load_default()
        small = ImageFont.load_default()

    nodes = {
        "@bdui/core": (80, 390),
        "@bdui/expr": (360, 170),
        "@bdui/defs": (360, 390),
        "@bdui/schema": (360, 610),
        "@bdui/dsl": (670, 170),
        "@bdui/runtime": (670, 390),
        "@bdui/transpiler": (670, 610),
        "@bdui/renderer-web": (1030, 80),
        "@bdui/registry": (1030, 290),
        "@bdui/sdk": (1030, 500),
        "@bdui/cli": (1030, 710),
    }
    edges = [
        ("@bdui/expr", "@bdui/core"),
        ("@bdui/defs", "@bdui/core"),
        ("@bdui/schema", "@bdui/core"),
        ("@bdui/dsl", "@bdui/core"),
        ("@bdui/dsl", "@bdui/defs"),
        ("@bdui/dsl", "@bdui/expr"),
        ("@bdui/runtime", "@bdui/core"),
        ("@bdui/runtime", "@bdui/expr"),
        ("@bdui/transpiler", "@bdui/core"),
        ("@bdui/transpiler", "@bdui/schema"),
        ("@bdui/renderer-web", "@bdui/core"),
        ("@bdui/renderer-web", "@bdui/defs"),
        ("@bdui/renderer-web", "@bdui/expr"),
        ("@bdui/renderer-web", "@bdui/runtime"),
        ("@bdui/registry", "@bdui/core"),
        ("@bdui/registry", "@bdui/schema"),
        ("@bdui/registry", "@bdui/runtime"),
        ("@bdui/sdk", "@bdui/core"),
        ("@bdui/sdk", "@bdui/schema"),
        ("@bdui/sdk", "@bdui/transpiler"),
        ("@bdui/sdk", "@bdui/registry"),
        ("@bdui/sdk", "@bdui/runtime"),
        ("@bdui/cli", "@bdui/core"),
        ("@bdui/cli", "@bdui/schema"),
        ("@bdui/cli", "@bdui/dsl"),
        ("@bdui/cli", "@bdui/transpiler"),
        ("@bdui/cli", "@bdui/registry"),
    ]
    box_w, box_h = 250, 72
    centers = {name: (x + box_w // 2, y + box_h // 2) for name, (x, y) in nodes.items()}
    for source, target in edges:
        sx, sy = centers[source]
        tx, ty = centers[target]
        draw_arrow(draw, (sx, sy), (tx, ty), "#9AA7B2")

    for label, (x, y) in nodes.items():
        draw.rounded_rectangle([x, y, x + box_w, y + box_h], radius=16, fill="#F5F8FC", outline="#2E5B89", width=3)
        bbox = draw.textbbox((0, 0), label, font=small)
        draw.text((x + (box_w - (bbox[2] - bbox[0])) / 2, y + 22), label, fill="#0B2545", font=small)

    draw.text((80, 35), "Граф зависимостей пакетов BDUI", fill="#0B2545", font=font)
    image.save(path)


def add_picture(doc: Document, path: Path):
    with Image.open(path) as image:
        width, height = image.size
    max_width = Inches(6.2)
    if height > width * 1.15:
        max_width = Inches(3.8)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run()
    run.add_picture(str(path), width=max_width)


def replace_page_count(text: str, page_count: int | None) -> str:
    if page_count is None:
        return text
    return re.sub(
        r"Работа содержит 4 рисунка, 0 таблиц, 34 листинга, 19 источников, 3 приложения; итоговое число страниц уточняется после верстки в шаблоне кафедры\.",
        f"Работа содержит {page_count} страниц, 4 рисунка, 0 таблиц, 34 листинга, 19 источников, 3 приложения.",
        text,
    )


def add_cover(doc: Document, lines: list[str]):
    for index, line in enumerate(lines):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_after = Pt(10 if index in (0, 3, 5) else 4)
        paragraph.paragraph_format.line_spacing = 1.15
        add_rich_text(paragraph, line, size=14, bold=index in (1, 5, 7))
    doc.add_page_break()


def add_toc(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    add_field(paragraph, r'TOC \o "1-3" \h \z \u', "Оглавление обновляется автоматически в Word.")
    doc.add_page_break()


def build_docx(source: Path, out: Path, page_count: int | None):
    create_dependency_graph(GRAPH_OUT)
    markdown = replace_page_count(source.read_text(encoding="utf-8"), page_count)
    lines = markdown.splitlines()

    doc = Document()
    configure_styles(doc)
    add_page_number_footer(doc)

    cover: list[str] = []
    i = 0
    while i < len(lines) and not lines[i].startswith("## "):
        if lines[i].strip():
            cover.append(lines[i].strip())
        i += 1
    add_cover(doc, cover)

    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    skip_manual_toc = False
    pending_paragraph: list[str] = []

    def flush_paragraph():
        nonlocal pending_paragraph
        if pending_paragraph:
            text = " ".join(item.strip() for item in pending_paragraph if item.strip())
            if text:
                if re.match(r"^(Рисунок|Листинг)\s+\d", text):
                    add_caption(doc, text)
                elif re.match(r"^\d+\.\s", text):
                    paragraph = doc.add_paragraph()
                    paragraph.paragraph_format.first_line_indent = None
                    paragraph.paragraph_format.left_indent = Cm(0.7)
                    paragraph.paragraph_format.first_line_indent = Cm(-0.7)
                    set_paragraph_spacing(paragraph, first_indent=False)
                    add_rich_text(paragraph, text)
                else:
                    add_body_paragraph(doc, text)
            pending_paragraph = []

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if in_code:
            if line.startswith("```"):
                flush_paragraph()
                if code_lang == "mermaid":
                    add_picture(doc, GRAPH_OUT)
                else:
                    add_code_block(doc, code_lines)
                in_code = False
                code_lang = ""
                code_lines = []
            else:
                code_lines.append(raw.rstrip())
            i += 1
            continue

        if line.startswith("```"):
            flush_paragraph()
            in_code = True
            code_lang = line.removeprefix("```").split(":", 1)[0].strip()
            code_lines = []
            i += 1
            continue

        if not line:
            flush_paragraph()
            i += 1
            continue

        if skip_manual_toc:
            if line.startswith("## ") and not line.startswith("## СОДЕРЖАНИЕ"):
                skip_manual_toc = False
            else:
                i += 1
                continue

        image_match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line)
        if image_match:
            flush_paragraph()
            add_picture(doc, (source.parent / image_match.group(1)).resolve())
            i += 1
            continue

        if line.startswith("### "):
            flush_paragraph()
            add_heading(doc, line[4:], 2)
            i += 1
            continue

        if line.startswith("## "):
            flush_paragraph()
            heading = line[3:]
            if heading == "СОДЕРЖАНИЕ":
                add_heading(doc, heading, 1)
                add_toc(doc)
                skip_manual_toc = True
            else:
                if len(doc.paragraphs) > 0:
                    doc.add_page_break()
                add_heading(doc, heading, 1)
            i += 1
            continue

        bullet_match = re.match(r"^[\-–]\s+(.*)", line)
        if bullet_match:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.first_line_indent = None
            add_rich_text(paragraph, bullet_match.group(1))
            i += 1
            continue

        pending_paragraph.append(line)
        i += 1

    flush_paragraph()

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def main():
    parser = argparse.ArgumentParser(description="Build the BDUI diploma DOCX from docs/vkr-part-2.md.")
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    parser.add_argument("--page-count", type=int, default=None)
    args = parser.parse_args()
    build_docx(args.source, args.out, args.page_count)
    print(args.out)


if __name__ == "__main__":
    main()
